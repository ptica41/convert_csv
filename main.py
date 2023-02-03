import copy

from reportlab.lib.pagesizes import landscape, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet
from deep_translator import GoogleTranslator
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor

import csv
import time


def parser(file='test.csv'):  # запись данных в список
    data = []

    with open(file, newline='') as f:
        reader = csv.reader(f)
        for row in reader:
            data.append(row)
    print('Количество строк:', len(data))

    return data


def edit(data: list):  # функция приведения списка к надлежащему виду

    temp = []
    clean = []
    sort = []
    fin = []
    risk = ['Risk Factor', 'Critical', 'High', 'Medium', 'Low']

    [temp.append(data[i]) for i in range(len(data)) if data[i][5] in risk]  # фильтрация по риск-фактору
    print('Количество строк после фильтрации по риск-фактору:', len(temp))

    for i in range(len(temp)):  # объединение колонок хост и порт
        temp[i].append(f'{temp[i][0]}:{temp[i][1]}')
        temp[i].pop(0)
        temp[i].pop(0)

    [clean.append(temp[i]) for i in range(len(temp)) if temp[i] not in clean]  # удаление дублируемых строк
    print('Количество строк после удаления дубликатов:', len(clean))

    for i in range(len(clean)):  # перестановка столбцов
        clean[i].insert(0, clean[i][3])
        clean[i].pop(4)

    for name in risk:  # сортировка по риск-фактору
        for i in range(len(clean)):
            if name == clean[i][0]:
                sort.append(clean[i])

    for i in range(len(sort)):  # объединение строк с одинаковым названием
        for j in range(len(sort)):
            if sort[i][1] == sort[j][1] and i != j:
                sort[i][4] = f'{sort[i][4]}\n{sort[j][4]}'
                sort[j][1] += 'na_uda1enie'

    for i in range(len(sort)):  # объединение строк с одинаковым названием
        if 'na_uda1enie' not in sort[i][1]:
            fin.append(sort[i])
    print('Количество строк после объединения:', len(fin))

    while True:  # Удаление лишних пробелов
        a = 0
        for i in range(len(fin)):
            for j in range(len(fin[i])):
                if fin[i][j].find('  ') != -1:
                    fin[i][j] = fin[i][j].replace('  ', ' ')
                    a += 1
        if not a:
            break

    while True:  # разбиение строк для переводчика (<5000 символов) и для полного отображения на странице
        count = 0  # счетчик количества новых строк после разбиения ячеек до 3000 символов
        for i in range(len(fin)):
            if len(fin[i][2]) > 2800:
                index = fin[i][2].rindex('\n\n', 0, 2800)
                fin.insert(i + 1, [fin[i][0], '', fin[i][2][index + 1:], '', ''])
                fin[i][2] = fin[i][2][:index]
                count += 1
        if not count:
            break

    return fin


def translate(data: list):  # перевод списка с английского на русский

    for i in range(1, len(data)):
        for j in range(len(data[i])):
            data[i][j] = GoogleTranslator(source='en', target='ru').translate(data[i][j])

    data[0] = ['УРОВЕНЬ РИСКА', 'УЯЗВИМОСТЬ ИЛИ НЕДОСТАТОК МЕХАНИЗМА ЗАЩИТЫ', 'ОПИСАНИЕ', 'РЕКОМЕНДАЦИИ',
               'УЯЗВИМЫЕ РЕСУРСЫ']  # замена "шапки"

    return data


def to_pdf(data: list, name='test.pdf'):  # генерация pdf-файла
    critical = 0
    high = 0
    medium = 0
    risk = ['Критический', 'Высокий', 'Средний', 'Низкий', 'Critical', 'High', 'Medium', 'Low']

    for i in range(len(data)):  # Удаление лишних переносов (кроме правого столбца) и исправление ошибки по тэгам
        for j in range(1, len(data[i]) - 1):
            data[i][j] = data[i][j].replace('<', ' < ')
            data[i][j] = data[i][j].replace('\n\n\n', '\n\n')
            data[i][j] = data[i][j].replace('\n\n', '<br />')
            data[i][j] = data[i][j].replace('\n', ' ')

    for i in range(len(data)):  # Удаление лишних пробелов
        for j in range(1, len(data[i]) - 1):
            data[i][j] = data[i][j].replace('  ', ' ')

    for i in range(len(data)):  # Замена переноса на тег (pdf не воспринимает \n) и коррекция перевода
        data[i][4] = data[i][4].replace('\n', '<br />')
        data[i][0] = data[i][0].replace('Середина', 'Средний')

    for i in range(1, len(data)):  # Счетчики количества строк по риск-фактору
        if data[i][0] == risk[0] or data[i][0] == risk[4]:
            critical += 1
        elif data[i][0] == risk[1] or data[i][0] == risk[5]:
            high += 1
        elif data[i][0] == risk[2] or data[i][0] == risk[6]:
            medium += 1

    pdfmetrics.registerFont(
        TTFont('arial', 'arial.ttf'))  # при необходимости замены шрифта нужно добавить его в папку проекта
    pdfmetrics.registerFont(TTFont('arialbd', 'arialbd.ttf'))

    stylesHead = getSampleStyleSheet()
    styleHead = stylesHead['BodyText']
    styleHead.fontName = 'arialbd'
    styleHead.textColor = '#5a5a5a'
    styleHead.alignment = 1  # центрирует текст внутри параграфа
    styleHead.fontSize = 8

    stylesName = getSampleStyleSheet()
    styleName = stylesName['BodyText']
    styleName.fontName = 'arialbd'
    styleName.textColor = '#000000'
    styleName.alignment = 1  # центрирует текст внутри параграфа
    styleName.fontSize = 8

    stylesText = getSampleStyleSheet()
    styleText = stylesText['BodyText']
    styleText.fontName = 'arial'
    styleText.textColor = '#000000'
    styleText.alignment = 1  # центрирует текст внутри параграфа
    styleText.fontSize = 7

    stylesRisk = getSampleStyleSheet()
    styleRisk = stylesRisk['BodyText']
    styleRisk.fontName = 'arialbd'
    styleRisk.textColor = '#ffffff'
    styleRisk.alignment = 1  # центрирует текст внутри параграфа
    styleRisk.fontSize = 8

    for i in range(len(data[0])):
        data[0][i] = Paragraph(data[0][i], styleHead)

    for i in range(1, len(data)):
        data[i][1] = Paragraph(data[i][1], styleName)

    for i in range(1, len(data)):
        for j in range(2, len(data[i])):
            data[i][j] = Paragraph(data[i][j], styleText)

    for i in range(1, len(data)):
        data[i][0] = Paragraph(data[i][0], styleRisk)

    t = Table(data, colWidths=[25 * mm, 40 * mm, 130 * mm, 54 * mm, 32 * mm], repeatRows=1,
              rowHeights=[15 * mm] + [None] * (len(data) - 1))
    t.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), (0.7, 0.7, 0.7)),
                           ('BOX', (0, 0), (-1, 0), 2, (1, 1, 1)),
                           ('INNERGRID', (0, 0), (-1, 0), 2, (1, 1, 1)),
                           ('BACKGROUND', (1, 1), (1, -1), (0.85, 0.85, 0.85)),
                           ('BOX', (1, 1), (1, -1), 2, (1, 1, 1)),
                           ('INNERGRID', (1, 1), (1, -1), 2, (1, 1, 1)),
                           ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                           ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                           ('LINEAFTER', (2, 1), (-1, -1), 0.5, (0.7, 0.7, 0.7)),
                           ('LINEBELOW', (2, 1), (-1, -1), 0.5, (0.7, 0.7, 0.7)),
                           ('BACKGROUND', (0, 1), (0, critical), '#670101'),
                           ('BACKGROUND', (0, critical + 1), (0, critical + high), '#970101'),
                           ('BACKGROUND', (0, critical + high + 1), (0, critical + high + medium), '#b29700'),
                           ('BACKGROUND', (0, critical + high + medium + 1), (0, -1), '#3f9500'),
                           ('BOX', (0, 1), (0, -1), 2, (1, 1, 1)),
                           ('INNERGRID', (0, 1), (0, -1), 2, (1, 1, 1)), ]))

    e = []
    e.append(t)
    doc = SimpleDocTemplate(name, pagesize=landscape(A4), rightMargin=10 * mm, leftMargin=10 * mm, topMargin=6 * mm,
                            bottomMargin=6 * mm, title='Convert csv to pdf', author='41')
    doc.build(e)


def to_docx(data: list, name='test.docx'):  # генерация pdf-файла
    critical = 0
    high = 0
    medium = 0
    risk = ['Критический', 'Высокий', 'Средний', 'Низкий', 'Critical', 'High', 'Medium', 'Low']

    for i in range(len(data)):  # Удаление лишних переносов для docx и коррекция перевода
        data[i][0] = data[i][0].replace('Середина', 'Средний')
        for j in range(1, len(data[i]) - 1):
            data[i][j] = data[i][j].replace('\n\n', '<br />')
            data[i][j] = data[i][j].replace('\n', ' ')
            data[i][j] = data[i][j].replace('<br />', '\n')

    for i in range(1, len(data)):  # Счетчики количества строк по риск-фактору
        if data[i][0] == risk[0] or data[i][0] == risk[4]:
            critical += 1
        elif data[i][0] == risk[1] or data[i][0] == risk[5]:
            high += 1
        elif data[i][0] == risk[2] or data[i][0] == risk[6]:
            medium += 1

    document = Document()
    section = document.sections[0]
    # section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    table = document.add_table(rows=len(data), cols=len(data[1]))
    table.autofit = False  # обязательный параметр для ручной настройки ширины колонок таблицы
    table.allow_autofit = False  # обязательный параметр для ручной настройки ширины колонок таблицы

    widths = (Cm(2.5), Cm(4), Cm(13), Cm(5.4), Cm(3.2))
    for row in table.rows:  # настройка ширины колонок таблицы
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    def set_repeat_table_header(row):  # для повтора шапки таблицы
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def preventDocumentBreak(document):  # запрет разделения строки на несколько страниц
        tags = document.element.xpath('//w:tr')
        rows = len(tags)
        for row in range(0, rows):
            tag = tags[row]
            child = OxmlElement('w:cantSplit')
            tag.append(child)

    def set_color_cell(cell, color_cell=None):  # цвет ячейки
        tblCell = cell._tc
        tblCellProperties = tblCell.get_or_add_tcPr()
        if color_cell:
            clShading = OxmlElement('w:shd')
            clShading.set(qn('w:fill'), color_cell)
            tblCellProperties.append(clShading)
        return cell

    def set_cell_border(cell, **kwargs):  # настройка границ ячейки
        """
        Пример использования:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    for i in range(len(data[0])):
        cell = table.cell(0, i)
        cell.text = data[0][i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.runs
        font = run[0].font
        font.name = 'Arial'
        font.bold = True
        font.color.rgb = RGBColor(0x5a, 0x5a, 0x5a)
        font.size = Pt(8)
        set_color_cell(cell, color_cell="b4b4b4")
        set_cell_border(cell, top={"sz": 24, "val": "single", "color": "#ffffff"},
                        bottom={"sz": 24, "val": "single", "color": "#ffffff"},
                        start={"sz": 24, "val": "single", "color": "#ffffff"},
                        end={"sz": 24, "val": "single", "color": "#ffffff"})

    for i in range(1, len(data)):
        cell = table.cell(i, 0)
        cell.text = data[i][0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.runs
        font = run[0].font
        font.name = 'Arial'
        font.bold = True
        font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        font.size = Pt(8)
        if i <= critical:
            set_color_cell(cell, color_cell="670101")
        elif i <= critical + high:
            set_color_cell(cell, color_cell="970101")
        elif i <= critical + high + medium:
            set_color_cell(cell, color_cell="b29700")
        else:
            set_color_cell(cell, color_cell="3f9500")
        set_cell_border(cell, top={"sz": 24, "val": "single", "color": "#ffffff"},
                        bottom={"sz": 24, "val": "single", "color": "#ffffff"},
                        start={"sz": 24, "val": "single", "color": "#ffffff"},
                        end={"sz": 24, "val": "single", "color": "#ffffff"})

    for i in range(1, len(data)):
        cell = table.cell(i, 1)
        cell.text = data[i][1]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.runs
        font = run[0].font
        font.name = 'Arial'
        font.bold = True
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        font.size = Pt(8)
        set_color_cell(cell, color_cell="dadada")
        set_cell_border(cell, top={"sz": 24, "val": "single", "color": "#ffffff"},
                        bottom={"sz": 24, "val": "single", "color": "#ffffff"},
                        start={"sz": 24, "val": "single", "color": "#ffffff"},
                        end={"sz": 24, "val": "single", "color": "#ffffff"})

    for i in range(1, len(data)):
        for j in range(2, len(data[i])):
            cell = table.cell(i, j)
            cell.text = data[i][j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.runs
            font = run[0].font
            font.name = 'Arial'
            font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            font.size = Pt(7)
            set_cell_border(cell, bottom={"sz": 6, "val": "single", "color": "#dadada"},
                            end={"sz": 6, "val": "single", "color": "#dadada"})

    set_repeat_table_header(table.rows[0])
    preventDocumentBreak(document)
    document.save(name)


if __name__ == '__main__':
    a = time.time()
    data = parser()  # при необходимости в аргументе функции указать название файла csv
    data = edit(data)
    data = translate(data)
    data_docx = copy.deepcopy(data)
    to_pdf(data)  # при необходимости в аргументе функции указать название файла pdf
    # to_docx(data)  # при необходимости в аргументе функции указать название файла docx
    b = (time.time() - a) / 60
    print('Время выполнения скрипта: %2.1f min' % b)
